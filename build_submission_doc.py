from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path('/Users/amit/pilates-landing')
OUT = ROOT / 'LUMERA_הגשה_פרסום_בגוגל_עמית_סולימן.docx'
LOGO = ROOT / 'assets/lumera-logo-landscape.png'

NAVY = '1F2A24'
SAGE = '5F7D68'
GOLD = 'C9A45C'
CREAM = 'F8F3EA'
MUTED = '6F706A'
WHITE = 'FFFFFF'


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def cell_border(cell, **kwargs):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge); element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag); borders.append(element)
            for key in ['val','sz','space','color']:
                if key in kwargs[edge]: element.set(qn('w:'+key), str(kwargs[edge][key]))


def set_cell_text(cell, text, bold=False, color=NAVY, size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.name = 'Arial'; r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def rtl_paragraph(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr(); bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)


def add_para(doc, text='', style=None, size=11, color=NAVY, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=7):
    p = doc.add_paragraph(style=style); rtl_paragraph(p); p.alignment = align
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.line_spacing = 1.15
    if text:
        r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); r.bold=bold
    return p


def heading(doc, text, level=1):
    p=doc.add_paragraph(); rtl_paragraph(p); p.paragraph_format.space_before=Pt(15 if level==1 else 9); p.paragraph_format.space_after=Pt(7)
    r=p.add_run(text); r.font.name='Arial'; r.font.bold=True; r.font.size=Pt(18 if level==1 else 13); r.font.color.rgb=RGBColor.from_string(SAGE if level==1 else NAVY)
    return p


def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); rtl_paragraph(p); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(NAVY)
    return p


def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,SAGE); set_cell_text(c,h,True,WHITE,10.5)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            shade(cells[i], CREAM if ri%2==0 else WHITE); set_cell_text(cells[i],str(val),False,NAVY,10)
    for row in t.rows:
        for c in row.cells: cell_border(c, top={'val':'single','sz':'5','color':'D8D1C3'}, bottom={'val':'single','sz':'5','color':'D8D1C3'}, left={'val':'single','sz':'5','color':'D8D1C3'}, right={'val':'single','sz':'5','color':'D8D1C3'})
    return t


doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.7); sec.right_margin=Inches(.7)
# default
styles=doc.styles
styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10.5); styles['Normal'].font.color.rgb=RGBColor.from_string(NAVY)
# header/footer
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.CENTER
hr=header.add_run(); hr.add_picture(str(LOGO), width=Inches(1.45))
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=footer.add_run('LUMERA  ·  פרסום בגוגל  ·  עמית סולימן'); fr.font.name='Arial'; fr.font.size=Pt(8); fr.font.color.rgb=RGBColor.from_string(MUTED)

# cover
p=add_para(doc,'',space_after=8); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(LOGO), width=Inches(3.7))
add_para(doc,'פרסום בגוגל',size=27,color=NAVY,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,space_after=2)
add_para(doc,'מטלת סיום · סמסטר ב׳ · תשפ״ו 2026',size=13,color=SAGE,align=WD_ALIGN_PARAGRAPH.CENTER,space_after=18)
# rule
p=add_para(doc,'',space_after=16); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━'); r.font.color.rgb=RGBColor.from_string(GOLD); r.font.size=Pt(12)
add_para(doc,'LUMERA',size=20,color=GOLD,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,space_after=2)
add_para(doc,'סטודיו לפילאטיס בוטיק בנתניה',size=14,color=MUTED,align=WD_ALIGN_PARAGRAPH.CENTER,space_after=25)
cover=table(doc,['פרט','מידע'],[
 ('סטודנט','עמית סולימן'),('תעודת זהות','213686108'),('טלפון','0587477769'),('אימייל','amitsuliman12@gmail.com'),('מרצה','אבישי פרוינד'),('קורס','פרסום בגוגל'),('מספר חשבון Google Ads','580-214-1692'),('תאריך הגשה','12/07/2026')])
add_para(doc,'',space_after=1)
add_para(doc,'מסמך זה מציג את הארגון, דפי הנחיתה ומבנה הפעילות הפרסומית שנבנתה עבור LUMERA.',size=11,color=MUTED,align=WD_ALIGN_PARAGRAPH.CENTER,space_after=0)
doc.add_page_break()

heading(doc,'1. תיאור הארגון')
add_para(doc,'LUMERA הוא סטודיו בוטיק לפילאטיס בנתניה ובאזור השרון. הסטודיו מציע אימוני רפורמר, שיעורי פילאטיס רצפה ואימונים פרטיים. המותג פונה למתאמנים ומתאמנות שמחפשים מסגרת מקצועית, יחס אישי ואימון שניתן להתאים לרמה ולמטרה שלהם.',size=11)
add_para(doc,'האתר נבנה כדף נחיתה רב־עמודי. בכל דף קיים טופס השארת פרטים, והשלמת הטופס מעבירה את המשתמש לדף תודה. בדף התודה מופעל אירוע המרה של Google Ads.',size=11)
heading(doc,'2. המוצרים שנבחרו',2)
add_para(doc,'א. פילאטיס מכשירים / רפורמר',size=12,bold=True,color=SAGE,space_after=3)
add_para(doc,'אימון על מכשיר רפורמר המשלב התנגדות, שליטה ודיוק. המסלול מתאים למתחילים ולמתאמנים מנוסים, ויכול לשמש לחיזוק שרירי הליבה, לשיפור היציבה, להגדלת טווחי תנועה ולבניית הרגלי תנועה נכונים. העבודה בקבוצה קטנה מאפשרת למדריך לתת תיקונים והתאמות במהלך השיעור.',size=10.8)
add_para(doc,'ב. אימון פילאטיס פרטי 1 על 1',size=12,bold=True,color=SAGE,space_after=3)
add_para(doc,'מסלול אישי שבו תוכן האימון, הקצב והתרגילים נבנים סביב המטרות והצרכים של המתאמן. האימון מתאים למי שמעדיף ליווי צמוד, למי שחוזר לפעילות לאחר הפסקה או פציעה, ולמי שמעוניין בתוכנית גמישה וממוקדת יותר.',size=10.8)
heading(doc,'3. דפי הנחיתה')
urls=[('דף הבית','https://avitheuser1-ship-it.github.io/pilates-landing/'),('פילאטיס מכשירים / רפורמר','https://avitheuser1-ship-it.github.io/pilates-landing/product1.html'),('חוג פילאטיס רצפה','https://avitheuser1-ship-it.github.io/pilates-landing/product2.html'),('אימון פרטי 1 על 1','https://avitheuser1-ship-it.github.io/pilates-landing/product3.html'),('דף תודה','https://avitheuser1-ship-it.github.io/pilates-landing/thank-you.html')]
table(doc,['דף','כתובת'],urls)
add_para(doc,'האתר אינו כולל תפריט ניווט, טלפון או הפניות לרשתות חברתיות. בדפים משולב הלוגו של LUMERA, והטופס מוביל לדף תודה לאחר שליחה.',size=10.5,color=MUTED)
doc.add_page_break()

heading(doc,'4. מבנה חשבון Google Ads')
table(doc,['קמפיין','סוג','תקציב יומי','סטטוס'],[
 ('קמפיין חיפוש','Search','30 ₪','מושהה'),('קמפיין מותג','Search','50 ₪','מושהה'),('קמפיין GDN','Display','40 ₪','מושהה'),('קמפיין וידאו','Demand Gen / נכסי וידאו','30 ₪','מושהה'),('קמפיין התקשרות בלבד','Search','40 ₪','מושהה')])
add_para(doc,'הקמפיינים נבנו במבנה הנדרש למטלה ונשמרו במצב מושהה, כך שלא יפעלו בפועל בזמן ההגשה.',size=10.8)
heading(doc,'5. קמפיין החיפוש והקמפיין הממותג')
add_para(doc,'בקמפיין החיפוש קיימות שתי קבוצות מודעות, כאשר החלוקה ביניהן מבוססת על צורך או אופן פנייה שונים. נעשה שימוש במילות מפתח בהתאמות מדויקת וביטוי, ללא התאמה רחבה וללא מילות מפתח כלליות בעלות מילה אחת בלבד.',size=10.8)
add_para(doc,'קמפיין המותג מיועד לחיפושים הקשורים ל־LUMERA. במודעות שולבה התאמה דינמית של מילת המפתח (DKI), כדי להתאים את נוסח המודעה לחיפוש המשתמש.',size=10.8)
heading(doc,'6. קהלים ורימרקטינג')
table(doc,['שם הקהל','תנאי עיקרי','שימוש'],[
 ('גולשים למוצר הראשון שלא המירו','ביקור ב־product1.html ללא ביקור בדף התודה','קבוצת GDN-לא-המירו'),('ביצעו המרה למוצר הראשון','ביקור ב־thank-you.html','קבוצת GDN-אפסייל')])
add_para(doc,'הקהל הראשון מיועד לחזרה למשתמשים שביקרו בדף הרפורמר אך לא השלימו טופס. הקהל השני מיועד למשתמשים שהשלימו את הטופס, ולכן ניתן לפנות אליהם בהצעת המשך או אפסייל.',size=10.8)
heading(doc,'7. קמפיין GDN וקמפיין הווידאו')
add_para(doc,'בקמפיין GDN קיימות שתי קבוצות מודעות: קבוצה למבקרים שלא המירו וקבוצה למשתמשים שביצעו המרה. לכל קבוצה הוגדר קהל הרימרקטינג המתאים.',size=10.8)
add_para(doc,'קמפיין הווידאו כולל קבוצת מודעות אחת ושני סרטוני YouTube מתאימים. בממשק החדש של Google Ads הקמפיין מופיע כ־Demand Gen עם נכסי וידאו, ומצב הטירגוט האופטימלי הוגדר כמושבת.',size=10.8)
add_para(doc,'סרטוני YouTube שנבחרו:',size=10.8,bold=True)
bullet(doc,'Pilates Reformer Workout | Full Body | Intermediate — Dez Fit')
bullet(doc,'Full Body Reformer Pilates Workout 40 Minutes — Alysia Pope Pilates')
heading(doc,'8. נכסים ותוספים')
table(doc,['נכס / תוסף','יישום / מצב'],[
 ('Sitelinks','קישורים לדף יצירת קשר ולדפי המוצרים — נכללו בייצוא'),('Callouts','הדגשת יתרונות, יחס אישי וקבוצות קטנות — נכללו בייצוא'),('Structured snippets','הצגת סוגי השירותים — נכללו בייצוא'),('Call asset','מספר הטלפון 0525381648, ישראל — קיים וממתין לבדיקה'),('Countdown','ספירה לאחור לקידום הצעה מוגבלת בזמן — הוקם בחשבון'),('מילים שליליות','סינון חיפושים שאינם רלוונטיים')])
heading(doc,'9. מדידה והמרות')
add_para(doc,'באתר משולב Google tag עם מזהה המדידה AW-18101301732. בדף thank-you.html מופעל אירוע ההמרה:',size=10.8)
add_para(doc,"AW-18101301732/7DiCCKiD0s4cEOTjr7dD",size=11,bold=True,color=GOLD,align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,'ההמרה נספרת לאחר שליחת הטופס והגעה לדף התודה. כך ניתן למדוד לא רק כניסות לדף, אלא גם פעולה עסקית של השארת פרטים.',size=10.8)
heading(doc,'10. סיכום')
add_para(doc,'העבודה מציגה תהליך של הקמת פעילות פרסום עבור עסק שירותי: אתר עם דפי נחיתה וטופס, מדידת המרות, קמפיין חיפוש וקמפיין מותג, קמפיין מדיה עם רימרקטינג, קמפיין וידאו ומודעות התקשרות. כל הקמפיינים בחשבון נשמרו במצב מושהה בהתאם להנחיות המטלה. בחשבון הוקמו גם תוסף ספירה לאחור, מודעת התקשרות בלבד, ארבעה מיקומי ערוצי YouTube ושני סרטוני YouTube.',size=11)

# final info box
p=doc.add_paragraph(); rtl_paragraph(p); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(0)
r=p.add_run('פרטי קשר להגשה: '); r.bold=True; r.font.color.rgb=RGBColor.from_string(SAGE); r.font.name='Arial'; r.font.size=Pt(10.5)
r=p.add_run('עמית סולימן · 0587477769 · amitsuliman12@gmail.com'); r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(NAVY)

doc.save(OUT)
print(OUT)
