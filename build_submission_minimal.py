from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path('/Users/amit/pilates-landing')
OUT = ROOT / 'LUMERA_הגשה_פרסום_בגוגל_עמית_סולימן.docx'
LOGO = ROOT / 'assets/lumera-logo-landscape-transparent.png'
SAGE='5F7D68'; GOLD='C9A45C'; CREAM='F8F3EA'; INK='1F2A24'; MUTED='6F706A'; WHITE='FFFFFF'

def bidi(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr=p._p.get_or_add_pPr(); e=OxmlElement('w:bidi'); e.set(qn('w:val'),'1'); pPr.append(e)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:shd'); e.set(qn('w:fill'),fill); tcPr.append(e)

def set_cell(cell, text, bold=False, color=INK, size=10.5):
    cell.text=''; p=cell.paragraphs[0]; bidi(p); p.paragraph_format.space_after=Pt(0)
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); r.bold=bold
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_p(doc,text='',size=11,color=INK,bold=False,align=WD_ALIGN_PARAGRAPH.RIGHT,after=7):
    p=doc.add_paragraph(); bidi(p); p.alignment=align; p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15
    if text:
        r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); r.bold=bold
    return p

def title(doc,text,size=17,color=SAGE):
    p=add_p(doc,text,size,color,True,after=7); p.paragraph_format.space_before=Pt(14); return p

def table(doc,headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers): shade(t.rows[0].cells[i],SAGE); set_cell(t.rows[0].cells[i],h,True,WHITE)
    for n,row in enumerate(rows):
        cs=t.add_row().cells
        for i,v in enumerate(row): shade(cs[i],CREAM if n%2==0 else WHITE); set_cell(cs[i],str(v))
    return t

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.6); sec.bottom_margin=Inches(.6); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(11)
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.CENTER; header.add_run().add_picture(str(LOGO),width=Inches(1.55))
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('LUMERA · פרסום בגוגל · עמית סולימן'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)

p=add_p(doc,'',after=3); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(LOGO),width=Inches(3.4))
add_p(doc,'פרסום בגוגל',25,INK,True,WD_ALIGN_PARAGRAPH.CENTER,2)
add_p(doc,'מטלת סיום · תשפ״ו 2026',13,SAGE,False,WD_ALIGN_PARAGRAPH.CENTER,20)

title(doc,'פרטי הסטודנט')
table(doc,['פרט','מידע'],[
 ('שם','עמית סולימן'),('תעודת זהות','213686108'),('טלפון','0587477769'),('אימייל','amitsuliman12@gmail.com')])

title(doc,'המוצרים')
table(doc,['מוצר','תיאור'],[
 ('פילאטיס מכשירים / רפורמר','אימון על מכשיר רפורמר לחיזוק, שיפור יציבה, שליטה בתנועה והתאמה לרמת המתאמן.'),
 ('אימון פילאטיס פרטי 1 על 1','אימון אישי עם תוכנית מותאמת, ליווי צמוד וגמישות בתוכן ובקצב האימון.')])

title(doc,'תיאור הארגון')
add_p(doc,'LUMERA הוא סטודיו בוטיק לפילאטיס בנתניה ובאזור השרון. הסטודיו מציע אימוני רפורמר ואימונים פרטיים, באווירה מקצועית ושקטה, עם דגש על יחס אישי, התאמה למטרות המתאמן ושיפור הכושר, היציבה ואיכות התנועה.',11,INK,False,after=10)

title(doc,'כתובות דפי הנחיתה')
table(doc,['דף','כתובת'],[
 ('דף הבית','https://avitheuser1-ship-it.github.io/pilates-landing/'),
 ('פילאטיס מכשירים / רפורמר','https://avitheuser1-ship-it.github.io/pilates-landing/product1.html'),
 ('חוג פילאטיס רצפה','https://avitheuser1-ship-it.github.io/pilates-landing/product2.html'),
 ('אימון פרטי 1 על 1','https://avitheuser1-ship-it.github.io/pilates-landing/product3.html'),
 ('דף תודה','https://avitheuser1-ship-it.github.io/pilates-landing/thank-you.html')])

title(doc,'מספר חשבון פרסום בגוגל')
add_p(doc,'580-214-1692',16,GOLD,True,WD_ALIGN_PARAGRAPH.CENTER,0)

doc.save(OUT); print(OUT)
